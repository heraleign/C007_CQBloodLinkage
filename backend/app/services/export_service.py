"""Lineage export service — Capability 6 (§9).

Generates Excel and Word documents for lineage analysis results.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.lineage_edge import LineageEdge
from app.models.lineage_node import LineageNode
from app.schemas.lineage import LineageExportRequest


class LineageExportService:
    """Generate lineage export documents."""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def export_excel(self, request: LineageExportRequest) -> tuple[bytes, str]:
        """Generate Excel export."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill

        wb = Workbook()
        filename = f"lineage_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        # Sheet 1: 表级血缘
        ws1 = wb.active
        ws1.title = "表级血缘"
        ws1.append(["源表名", "源系统", "目标表名", "目标系统", "血缘阶段", "采集方式", "置信度"])
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="E6F7FF", end_color="E6F7FF", fill_type="solid")
        for cell in ws1[1]:
            cell.font = header_font
            cell.fill = header_fill

        for table_name in request.table_names:
            nodes, edges = await self._get_lineage_data(table_name, request)
            for edge in edges:
                source_node = next((n for n in nodes if n.node_id == edge.source_node_id), None)
                target_node = next((n for n in nodes if n.node_id == edge.target_node_id), None)
                ws1.append([
                    source_node.table_name if source_node else edge.source_node_id,
                    source_node.system_name if source_node else "",
                    target_node.table_name if target_node else edge.target_node_id,
                    target_node.system_name if target_node else "",
                    edge.lineage_stage,
                    edge.collect_method or "",
                    float(edge.confidence) if edge.confidence else 1.0,
                ])

        # Sheet 2: 节点详情
        ws2 = wb.create_sheet("节点详情")
        ws2.append(["节点名", "类型", "归属系统", "集群", "数据库", "表名", "备注"])
        for cell in ws2[1]:
            cell.font = header_font
            cell.fill = header_fill

        all_nodes_set = set()
        for table_name in request.table_names:
            nodes, _ = await self._get_lineage_data(table_name, request)
            for n in nodes:
                if n.node_id not in all_nodes_set:
                    all_nodes_set.add(n.node_id)
                    ws2.append([
                        n.node_name, n.node_type, n.system_name or "",
                        n.cluster_name or "", n.database_name or "",
                        n.table_name or "", n.table_remark or "",
                    ])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.getvalue(), filename

    async def export_word(self, request: LineageExportRequest) -> tuple[bytes, str]:
        """Generate Word document export."""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        filename = f"lineage_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Title
        title = doc.add_heading("血缘分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Overview
        doc.add_heading("1. 血缘分析概述", level=1)
        doc.add_paragraph(f"分析对象：{', '.join(request.table_names)}")
        doc.add_paragraph(f"分析时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"分析方向：{'全部' if request.direction == 'BOTH' else request.direction}")
        doc.add_paragraph(f"分析层数：{request.depth}")

        # Table-level lineage
        doc.add_heading("2. 表级血缘明细", level=1)
        for table_name in request.table_names:
            doc.add_heading(f"2.1 {table_name}", level=2)
            nodes, edges = await self._get_lineage_data(table_name, request)

            if edges:
                table_data = []
                table_data.append(["源表", "目标表", "阶段", "采集方式"])
                for edge in edges:
                    source_node = next((n for n in nodes if n.node_id == edge.source_node_id), None)
                    target_node = next((n for n in nodes if n.node_id == edge.target_node_id), None)
                    table_data.append([
                        source_node.table_name if source_node else "",
                        target_node.table_name if target_node else "",
                        edge.lineage_stage,
                        edge.collect_method or "",
                    ])
                if len(table_data) > 1:
                    table = doc.add_table(rows=len(table_data), cols=4)
                    table.style = "Light Grid Accent 1"
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            table.cell(i, j).text = cell_text

        # Finalize
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), filename

    async def _get_lineage_data(
        self, table_name: str, request: LineageExportRequest
    ) -> tuple[list[Any], list[Any]]:
        """Get lineage nodes and edges for a table."""
        # Find the node
        stmt = select(LineageNode).where(
            LineageNode.table_name == table_name,
            LineageNode.status == 1,
        )
        result = await self.db.execute(stmt)
        center = result.scalar_one_or_none()
        if not center:
            return [], []

        # Get connected edges
        edge_stmt = select(LineageEdge).where(
            (LineageEdge.source_node_id == center.node_id) | (LineageEdge.target_node_id == center.node_id),
            LineageEdge.status == 1,
        )
        if request.lineage_type != "ALL":
            edge_stmt = edge_stmt.where(LineageEdge.lineage_type == request.lineage_type)
        edge_result = await self.db.execute(edge_stmt)
        edges = list(edge_result.scalars().all())

        # Get all connected nodes
        node_ids = {center.node_id}
        for edge in edges:
            node_ids.add(edge.source_node_id)
            node_ids.add(edge.target_node_id)

        node_stmt = select(LineageNode).where(LineageNode.node_id.in_(node_ids))
        node_result = await self.db.execute(node_stmt)
        nodes = list(node_result.scalars().all())

        return nodes, edges
